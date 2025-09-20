from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
import json
import os
from models import db, init_db, Prompt, SavedPrompt
from dotenv import load_dotenv
from docx import Document
from openpyxl import Workbook

load_dotenv()

app = Flask(__name__)
CORS(app)
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL', 'sqlite:///prompts.db')
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'default-secret')

db.init_app(app)

@app.route('/')
def index():
    with open('cats.json', 'r', encoding='utf-8') as f:
        cats = json.load(f)
    cats_json = json.dumps(cats)
    return render_template('index.html', cats_json=cats_json)

@app.route('/gerar_prompt', methods=['POST'])
def gerar_prompt():
    data = request.get_json()
    # Convert data types
    data['historico'] = data.get('historico') == 'on'
    data['modo_iterativo'] = data.get('modo_iterativo') == 'on'
    data['nivel_criatividade'] = int(data.get('nivel_criatividade', 50))
    # Convert lists to JSON strings for storage
    data['ferramentas_externas'] = json.dumps(data.get('ferramentas_externas', []))
    data['metricas'] = json.dumps(data.get('metricas', []))
    data['integracao_multimodal'] = json.dumps(data.get('integracao_multimodal', []))
    data['testes'] = json.dumps(data.get('testes', []))
    data['integracao_ci_cd'] = json.dumps(data.get('integracao_ci_cd', []))
    # Generate JSON based on data
    full_prompt_json = {
        "titulo": data.get('titulo'),
        "tipo": data.get('tipo'),
        "subtipo": data.get('subtipo'),
        "subitem": data.get('subitem'),
        "tom_estilo": data.get('tom_estilo'),
        "persona": data.get('persona'),
        "objetivos": data.get('objetivos'),
        "regras": data.get('regras'),
        "contexto": data.get('contexto'),
        "restricoes": data.get('restricoes'),
        "exemplos": data.get('exemplos'),
        "framework": data.get('framework'),
        "frontend": data.get('frontend'),
        "backend": data.get('backend'),
        "main_language": data.get('main_language'),
        "ferramentas_externas": json.loads(data.get('ferramentas_externas')),
        "nivel_detalhe": data.get('nivel_detalhe'),
        "publico_alvo": data.get('publico_alvo'),
        "formato_saida": data.get('formato_saida'),
        "prioridade": data.get('prioridade'),
        "metodo_raciocinio": data.get('metodo_raciocinio'),
        "idioma": data.get('idioma'),
        "metricas": json.loads(data.get('metricas')),
        "integracao_multimodal": json.loads(data.get('integracao_multimodal')),
        "historico": data.get('historico'),
        "modo_iterativo": data.get('modo_iterativo'),
        "nivel_criatividade": data.get('nivel_criatividade'),
        "dependencias": data.get('dependencias'),
        "ambiente_deploy": data.get('ambiente_deploy'),
        "testes": json.loads(data.get('testes')),
        "versionamento": data.get('versionamento'),
        "estrutura_projeto": data.get('estrutura_projeto'),
        "padroes_codigo": data.get('padroes_codigo'),
        "integracao_ci_cd": json.loads(data.get('integracao_ci_cd')),
        "estilo_saida": data.get('estilo_saida')
    }

    # Filter out empty values
    def is_empty(value):
        if value is None:
            return True
        if isinstance(value, str) and value.strip() == '':
            return True
        if isinstance(value, list) and len(value) == 0:
            return True
        if isinstance(value, bool) and not value:
            return True
        return False

    filtered_json = {k: v for k, v in full_prompt_json.items() if not is_empty(v)}

    # Adjust based on nivel_detalhe
    nivel = data.get('nivel_detalhe')
    if nivel == 'Resumido':
        required_keys = ["titulo", "tipo", "objetivos", "formato_saida"]
        prompt_json = {k: filtered_json[k] for k in required_keys if k in filtered_json}
    elif nivel == 'Intermediário':
        required_keys = ["titulo", "tipo", "subtipo", "objetivos", "regras", "contexto", "formato_saida"]
        prompt_json = {k: filtered_json[k] for k in required_keys if k in filtered_json}
    else:  # Detalhado
        prompt_json = filtered_json
    # Save to DB
    prompt = Prompt(**data)
    db.session.add(prompt)
    db.session.commit()
    return jsonify({"id": prompt.id, "data": prompt_json})

@app.route('/download/<int:prompt_id>/<format>')
def download(prompt_id, format):
    prompt = Prompt.query.get_or_404(prompt_id)
    data = {
        "titulo": prompt.titulo,
        "tipo": prompt.tipo,
        "subtipo": prompt.subtipo,
        "subitem": prompt.subitem,
        "tom_estilo": prompt.tom_estilo,
        "persona": prompt.persona,
        "objetivos": prompt.objetivos,
        "regras": prompt.regras,
        "contexto": prompt.contexto,
        "restricoes": prompt.restricoes,
        "exemplos": prompt.exemplos,
        "framework": prompt.framework,
        "frontend": prompt.frontend,
        "backend": prompt.backend,
        "main_language": prompt.main_language,
        "ferramentas_externas": json.loads(prompt.ferramentas_externas) if prompt.ferramentas_externas else [],
        "nivel_detalhe": prompt.nivel_detalhe,
        "publico_alvo": prompt.publico_alvo,
        "formato_saida": prompt.formato_saida,
        "prioridade": prompt.prioridade,
        "metodo_raciocinio": prompt.metodo_raciocinio,
        "idioma": prompt.idioma,
        "metricas": json.loads(prompt.metricas) if prompt.metricas else [],
        "integracao_multimodal": json.loads(prompt.integracao_multimodal) if prompt.integracao_multimodal else [],
        "historico": prompt.historico,
        "modo_iterativo": prompt.modo_iterativo,
        "nivel_criatividade": prompt.nivel_criatividade,
        "dependencias": prompt.dependencias,
        "ambiente_deploy": prompt.ambiente_deploy,
        "testes": json.loads(prompt.testes) if prompt.testes else [],
        "versionamento": prompt.versionamento,
        "estrutura_projeto": prompt.estrutura_projeto,
        "padroes_codigo": prompt.padroes_codigo,
        "integracao_ci_cd": json.loads(prompt.integracao_ci_cd) if prompt.integracao_ci_cd else [],
        "estilo_saida": prompt.estilo_saida
    }

    # Filter out empty values
    def is_empty(value):
        if value is None:
            return True
        if isinstance(value, str) and value.strip() == '':
            return True
        if isinstance(value, list) and len(value) == 0:
            return True
        if isinstance(value, bool) and not value:
            return True
        return False

    filtered_data = {k: v for k, v in data.items() if not is_empty(v)}

    if format == 'json':
        filename = f'prompt_{prompt_id}.json'
        with open(filename, 'w') as f:
            json.dump(filtered_data, f, indent=2)
        return send_file(filename, as_attachment=True)
    elif format == 'txt':
        filename = f'prompt_{prompt_id}.txt'
        with open(filename, 'w') as f:
            f.write(json.dumps(filtered_data, indent=2))
        return send_file(filename, as_attachment=True)
    elif format == 'css':
        filename = f'prompt_{prompt_id}.css'
        with open(filename, 'w') as f:
            f.write("/* Generated CSS */\n")
            f.write(json.dumps(filtered_data, indent=2))
        return send_file(filename, as_attachment=True)
    elif format == 'docx':
        filename = f'prompt_{prompt_id}.docx'
        doc = Document()
        doc.add_heading('Prompt Data', 0)
        for key, value in filtered_data.items():
            doc.add_paragraph(f'{key}: {value}')
        doc.save(filename)
        return send_file(filename, as_attachment=True)
    elif format == 'xlsx':
        filename = f'prompt_{prompt_id}.xlsx'
        wb = Workbook()
        ws = wb.active
        ws.title = "Prompt Data"
        row = 1
        for key, value in filtered_data.items():
            ws.cell(row=row, column=1, value=key)
            ws.cell(row=row, column=2, value=str(value))
            row += 1
        wb.save(filename)
        return send_file(filename, as_attachment=True)

@app.route('/save_prompt', methods=['POST'])
def save_prompt():
    data = request.get_json()
    name = data.get('name')
    prompt_data = data.get('data')
    if not name or not prompt_data:
        return jsonify({"error": "Nome e dados são obrigatórios"}), 400
    saved = SavedPrompt(name=name, data=json.dumps(prompt_data))
    db.session.add(saved)
    db.session.commit()
    return jsonify({"id": saved.id, "name": saved.name})

@app.route('/load_prompts')
def load_prompts():
    prompts = SavedPrompt.query.all()
    result = [{"id": p.id, "name": p.name} for p in prompts]
    return jsonify(result)

@app.route('/load_prompt/<int:prompt_id>')
def load_prompt(prompt_id):
    prompt = SavedPrompt.query.get_or_404(prompt_id)
    return jsonify(json.loads(prompt.data))

@app.route('/rename_prompt/<int:prompt_id>', methods=['POST'])
def rename_prompt(prompt_id):
    data = request.get_json()
    new_name = data.get('name')
    if not new_name:
        return jsonify({"error": "Nome é obrigatório"}), 400
    prompt = SavedPrompt.query.get_or_404(prompt_id)
    prompt.name = new_name
    db.session.commit()
    return jsonify({"success": True})

@app.route('/delete_prompt/<int:prompt_id>', methods=['DELETE'])
def delete_prompt(prompt_id):
    prompt = SavedPrompt.query.get_or_404(prompt_id)
    db.session.delete(prompt)
    db.session.commit()
    return jsonify({"success": True})

if __name__ == '__main__':
    init_db(app)
    app.run(debug=True)